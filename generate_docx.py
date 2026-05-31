import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.opc.constants
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import re


def add_hyperlink(paragraph, url, text):
    # Get access to the document.xml.rels file and get a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    new_run = OxmlElement("w:r")

    # Create a new w:rPr element
    rPr = OxmlElement("w:rPr")

    # Add color and underline to the run properties
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")  # Standard link blue
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Join elements
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(33, 56, 91)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(33, 56, 91)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)


def process_text_with_links(p, text):
    # Regex to find links [text](url) and bold **text** and italic *text*
    # We will split the text into a list of chunks: plain text, links, bolds, italics
    pattern = r"(\[.*?\]\(.*?\)|\*\*.*?\*\*|\*[^\*]+\*)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("[") and "](" in part and part.endswith(")"):
            # It's a link
            title, url = re.match(r"\[(.*?)\]\((.*?)\)", part).groups()
            add_hyperlink(p, url, title)
        elif part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)


def parse_markdown(md_file, docx_file):
    doc = docx.Document()
    set_style(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            process_text_with_links(p, line[2:])
        else:
            p = doc.add_paragraph()
            if "|" in line and "@" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            process_text_with_links(p, line)

    doc.save(docx_file)
    print("Done generating", docx_file)


if __name__ == "__main__":
    parse_markdown("Alexander_Tyutin_CV_Updated.md", "Alexander_Tyutin_CV.docx")
