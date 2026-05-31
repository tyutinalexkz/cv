"""
This module generates a Word document (DOCX) and a PDF from a Markdown CV.
"""

import re
import docx
import docx.opc.constants
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

try:
    from fpdf import FPDF
    import markdown

    HAS_PDF_TOOLS = True
except ImportError:
    HAS_PDF_TOOLS = False


def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a docx paragraph.
    """
    # Get access to the document.xml.rels file and get a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    new_run = OxmlElement("w:r")

    # Create a new w:rPr element
    r_pr = OxmlElement("w:rPr")

    # Add color and underline to the run properties
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")  # Standard link blue
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    # Join elements
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)

    # pylint: disable=protected-access
    paragraph._p.append(hyperlink)
    return hyperlink


def set_style(doc):
    """
    Configures the default styles (font, size, color) for the document.
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(33, 56, 91)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(33, 56, 91)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)


def process_text_with_links(p, text):
    """
    Processes markdown text with links and formatting, adding them to a paragraph.
    """
    # Regex to find links [text](url) and bold **text** and italic *text*
    # We will split the text into a list of chunks: plain text, links, bolds, italics
    pattern = r"(\[.*?\]\(.*?\)|\*\*.*?\*\*|\*[^\*]+\*)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("[") and "](" in part and part.endswith(")"):
            # It's a link
            title, url = re.match(r"\[(.*?)\]\((.*?)\)", part).groups()
            add_hyperlink(p, url, title)
        elif part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)


def parse_markdown(md_file, docx_file):
    """
    Reads a markdown file and generates a stylized docx document.
    """
    doc = docx.Document()
    set_style(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            process_text_with_links(p, line[2:])
        else:
            p = doc.add_paragraph()
            if "|" in line and "@" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            process_text_with_links(p, line)

    doc.save(docx_file)
    print("Done generating", docx_file)


def main():
    """
    Main execution function: parses the markdown and optionally generates a PDF.
    """
    docx_filename = "cv.docx"
    pdf_filename = "cv.pdf"
    parse_markdown("cv.md", docx_filename)

    if HAS_PDF_TOOLS:
        print(f"Generating {pdf_filename} from markdown...")
        with open("cv.md", "r", encoding="utf-8") as f:
            text = f.read()

        # Clean up unicode characters not supported by default FPDF font
        text = (
            text.replace("—", "-")
            .replace("’", "'")
            .replace("“", '"')
            .replace("”", '"')
            .replace("…", "...")
        )

        html = markdown.markdown(text)

        # Increase line spacing with inline CSS
        html = html.replace("<p>", "<p style='line-height: 1.2'>")
        # For bulleted lists, we wrap the content in a <p> with line-height and add a <br>
        html = html.replace("<li>", "<li style='line-height: 0.0'>").replace(
            "</li>", "<br></li>"
        )

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        pdf.write_html(html)
        pdf.output(pdf_filename)
        print("Done generating", pdf_filename)
    else:
        print("fpdf2 or markdown is not installed. Skipping PDF generation.")


if __name__ == "__main__":
    main()
